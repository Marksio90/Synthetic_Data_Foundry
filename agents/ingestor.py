"""
agents/ingestor.py — Agent Ingestor (Multi-format Parser)

Obsługiwane formaty:
  PDF  → LlamaParse (cloud, tabele) → PyMuPDF (local) → OpenAI Vision (skany)
  DOCX → python-docx (local, darmowy)
  HTML → BeautifulSoup (local, darmowy)
  XML  → BeautifulSoup (local, darmowy)
  TXT  → plain text (local, darmowy)
  MD   → Markdown as-is (local, darmowy)
  MP3/WAV/M4A/MP4 → Replicate Whisper large-v3 (transkrypcja audio → markdown)

Wymagania:
  python-docx   — pip install python-docx
  beautifulsoup4 — pip install beautifulsoup4 lxml
  replicate     — pip install replicate (do audio)
"""

from __future__ import annotations

import hashlib
import logging
import re
from datetime import date
from pathlib import Path
from typing import Optional

from sqlalchemy.orm import Session

from config.settings import settings
from db import repository as repo

logger = logging.getLogger(__name__)

# Obsługiwane rozszerzenia
_AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".mp4", ".ogg", ".flac", ".webm"}
_PDF_EXTENSIONS = {".pdf"}
_DOCX_EXTENSIONS = {".docx", ".doc"}
_HTML_EXTENSIONS = {".html", ".htm", ".xml"}
_TEXT_EXTENSIONS = {".txt", ".md", ".rst"}

_DATE_RE = re.compile(
    r"(?:wchodzi w życie|applicable from|valid from|effective|od dnia)\s*:?\s*"
    r"(\d{1,2}[\./\-]\d{1,2}[\./\-]\d{4}|\d{4}[\./\-]\d{2}[\./\-]\d{2})",
    re.IGNORECASE,
)
_YEAR_RE = re.compile(r"\b(20\d{2})\b")


# ---------------------------------------------------------------------------
# Parsery — PDF
# ---------------------------------------------------------------------------

def _parse_with_llamaparse(pdf_path: Path) -> str:
    from llama_parse import LlamaParse  # type: ignore

    parser = LlamaParse(
        api_key=settings.llama_parse_api_key,
        result_type="markdown",
        language="en",
        verbose=False,
    )
    documents = parser.load_data(str(pdf_path))
    return "\n\n".join(doc.text for doc in documents)


def _parse_with_pymupdf(pdf_path: Path) -> str:
    import fitz  # type: ignore

    doc = fitz.open(str(pdf_path))
    pages_md: list[str] = []

    for page_num, page in enumerate(doc, start=1):
        page_dict = page.get_text("dict")
        lines: list[str] = []

        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_parts: list[str] = []
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    font_size = span.get("size", 0)
                    if font_size > 13:
                        line_parts.append(f"## {text}")
                    else:
                        line_parts.append(text)
                if line_parts:
                    lines.append(" ".join(line_parts))

        page_md = "\n".join(lines)
        pages_md.append(f"<!-- Page {page_num} -->\n{page_md}")

    doc.close()
    return "\n\n".join(pages_md)


def _parse_with_openai_vision(pdf_path: Path) -> str:
    import base64
    import io

    import openai
    from pdf2image import convert_from_path  # type: ignore
    from PIL import Image  # type: ignore  # noqa: F401

    client = openai.OpenAI(api_key=settings.openai_api_key)
    pages = convert_from_path(str(pdf_path), dpi=150)
    full_markdown: list[str] = []

    for i, page in enumerate(pages):
        buf = io.BytesIO()
        page.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Convert this document page to Markdown. "
                            "Reproduce all tables using Markdown table syntax. "
                            "Preserve all article numbers, headings, and structure exactly."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{b64}"},
                    },
                ],
            }],
            max_tokens=4096,
        )
        full_markdown.append(f"<!-- Page {i+1} -->\n{response.choices[0].message.content}")

    return "\n\n".join(full_markdown)


# ---------------------------------------------------------------------------
# Parsery — DOCX
# ---------------------------------------------------------------------------

def _parse_with_docx(docx_path: Path) -> str:
    """Parse DOCX → Markdown using python-docx."""
    try:
        import docx  # type: ignore
    except ImportError:
        raise ImportError("python-docx not installed: pip install python-docx")

    doc = docx.Document(str(docx_path))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        else:
            lines.append(text)

    # Tables
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Parsery — HTML / XML
# ---------------------------------------------------------------------------

def _parse_with_html(html_path: Path) -> str:
    """Parse HTML/XML → Markdown using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        raise ImportError("beautifulsoup4 not installed: pip install beautifulsoup4 lxml")

    with html_path.open("r", encoding="utf-8", errors="replace") as f:
        content = f.read()

    soup = BeautifulSoup(content, "lxml")

    # Remove script/style
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    lines: list[str] = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "table", "tr"]):
        tag = element.name
        text = element.get_text(separator=" ").strip()
        if not text:
            continue
        if tag == "h1":
            lines.append(f"# {text}")
        elif tag == "h2":
            lines.append(f"## {text}")
        elif tag == "h3":
            lines.append(f"### {text}")
        elif tag == "h4":
            lines.append(f"#### {text}")
        elif tag == "li":
            lines.append(f"- {text}")
        elif tag in ("p", "tr"):
            lines.append(text)

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Parsery — Plain text / Markdown
# ---------------------------------------------------------------------------

def _parse_plain_text(path: Path) -> str:
    with path.open("r", encoding="utf-8", errors="replace") as f:
        return f.read()


# ---------------------------------------------------------------------------
# Parsery — Audio via Replicate Whisper large-v3
# ---------------------------------------------------------------------------

def _parse_with_whisper(audio_path: Path) -> str:
    """
    Transkrypcja audio → Markdown za pomocą Replicate Whisper large-v3.
    Koszt: ~$0.0002/min audio (nagrania konferencji ESG → dane treningowe).
    """
    if not settings.replicate_api_key:
        raise ValueError(
            "REPLICATE_API_KEY not set — audio transcription unavailable. "
            "Get your key at https://replicate.com"
        )

    try:
        import replicate  # type: ignore
    except ImportError:
        raise ImportError("replicate not installed: pip install replicate")

    import os
    os.environ.setdefault("REPLICATE_API_TOKEN", settings.replicate_api_key)

    logger.info("Transcribing audio with Whisper large-v3: %s", audio_path.name)

    with audio_path.open("rb") as audio_file:
        output = replicate.run(
            "openai/whisper:4d50797290df275329f202e48c76360b3f22b08d28c196cbc54600319435f8d2",
            input={
                "audio": audio_file,
                "model": "large-v3",
                "language": "auto",
                "transcription": "plain text",
                "word_timestamps": False,
            },
        )

    transcript = output.get("transcription", "") if isinstance(output, dict) else str(output)

    # Format jako Markdown z nagłówkiem
    return f"# Transkrypcja: {audio_path.name}\n\n{transcript}"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _extract_valid_from_date(markdown: str) -> Optional[date]:
    match = _DATE_RE.search(markdown)
    if not match:
        return None
    raw = match.group(1)
    for fmt in ("%d.%m.%Y", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%Y.%m.%d"):
        try:
            from datetime import datetime
            return datetime.strptime(raw, fmt).date()
        except ValueError:
            continue
    return None


def _extract_directive_year(markdown: str, filename: str) -> Optional[int]:
    for text in (filename, markdown[:500]):
        match = _YEAR_RE.search(text)
        if match:
            return int(match.group(1))
    return None


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


# ---------------------------------------------------------------------------
# Routing — wybór parsera na podstawie rozszerzenia
# ---------------------------------------------------------------------------

def _parse_document(path: Path) -> str:
    """Route file to appropriate parser based on extension."""
    ext = path.suffix.lower()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if ext in _PDF_EXTENSIONS:
        markdown = ""
        if settings.llama_parse_api_key:
            logger.info("Parsing PDF with LlamaParse: %s", path.name)
            try:
                markdown = _parse_with_llamaparse(path)
            except Exception as exc:
                logger.warning("LlamaParse failed (%s), falling back to PyMuPDF", exc)
        if not markdown:
            logger.info("Parsing PDF with PyMuPDF: %s", path.name)
            try:
                markdown = _parse_with_pymupdf(path)
            except Exception as exc:
                logger.warning("PyMuPDF failed (%s), falling back to OpenAI Vision", exc)
        if not markdown:
            logger.info("Parsing PDF with OpenAI Vision: %s", path.name)
            markdown = _parse_with_openai_vision(path)
        return markdown

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if ext in _DOCX_EXTENSIONS:
        logger.info("Parsing DOCX: %s", path.name)
        return _parse_with_docx(path)

    # ── HTML / XML ────────────────────────────────────────────────────────────
    if ext in _HTML_EXTENSIONS:
        logger.info("Parsing HTML/XML: %s", path.name)
        return _parse_with_html(path)

    # ── Audio ─────────────────────────────────────────────────────────────────
    if ext in _AUDIO_EXTENSIONS:
        logger.info("Parsing audio with Replicate Whisper: %s", path.name)
        return _parse_with_whisper(path)

    # ── Plain text / Markdown ─────────────────────────────────────────────────
    if ext in _TEXT_EXTENSIONS:
        logger.info("Reading text file: %s", path.name)
        return _parse_plain_text(path)

    # ── Unknown — try plain text as last resort ───────────────────────────────
    logger.warning("Unknown file type '%s' — treating as plain text: %s", ext, path.name)
    return _parse_plain_text(path)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def ingest_document(file_path: str, session: Session) -> tuple[str, str]:
    """
    Parse *file_path* (PDF/DOCX/HTML/TXT/audio) and persist a SourceDocument.

    Returns (source_doc_id, raw_markdown).
    Backward-compatible alias: ingest_pdf() still works for PDF files.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    file_hash = _sha256(path)

    # Dedup check
    from sqlalchemy import select
    from db.models import SourceDocument
    existing = session.scalar(
        select(SourceDocument).where(SourceDocument.file_hash == file_hash)
    )
    if existing:
        logger.info("Document already ingested: %s — skipping", path.name)
        return str(existing.id), existing.raw_markdown or ""

    # Parse
    markdown = _parse_document(path)

    valid_from = _extract_valid_from_date(markdown)
    directive_year = _extract_directive_year(markdown, path.stem)

    doc = repo.upsert_source_document(
        session,
        filename=path.name,
        file_hash=file_hash,
        directive_name=path.stem.upper(),
        directive_year=directive_year,
        valid_from_date=valid_from,
        raw_markdown=markdown,
    )
    session.commit()
    logger.info(
        "Ingested '%s' [%s] → doc_id=%s valid_from=%s chars=%d",
        path.name,
        path.suffix.upper(),
        doc.id,
        valid_from,
        len(markdown),
    )
    return str(doc.id), markdown


# Backward-compatible alias
def ingest_pdf(pdf_path: str, session: Session) -> tuple[str, str]:
    return ingest_document(pdf_path, session)
